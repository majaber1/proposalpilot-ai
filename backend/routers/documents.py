"""
Documents Router - Handles file upload and management
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel
from typing import Optional, List
import os
import uuid
import aiofiles
from pathlib import Path

from database import get_db, Document

router = APIRouter()

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/app/uploads")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE_MB", "50")) * 1024 * 1024
ALLOWED_TYPES = {"application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"}


async def extract_text(file_path: str, file_type: str) -> str:
    """Extract text from uploaded file"""
    try:
        if file_type == "text/plain" or file_path.endswith(".txt"):
            async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return await f.read()
        elif file_path.endswith(".pdf"):
            import PyPDF2
            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        elif file_path.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        return "Text extraction not supported for this file type"
    except Exception as e:
        return f"Error extracting text: {str(e)}"


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db)
):
    """Upload an RFP document"""
    # Validate file
    if file.size and file.size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large")
    
    # Generate unique filename
    file_ext = Path(file.filename).suffix.lower()
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    # Save file
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    content = await file.read()
    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)
    
    # Extract text
    extracted_text = await extract_text(file_path, file.content_type or "")
    
    # Save to database
    document = Document(
        user_id=1,  # Default user for MVP
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_type=file.content_type,
        file_size=len(content),
        status="analyzed" if extracted_text else "uploaded",
        extracted_text=extracted_text
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)
    
    return {
        "id": document.id,
        "filename": document.original_filename,
        "status": document.status,
        "text_length": len(extracted_text) if extracted_text else 0,
        "message": "Document uploaded and text extracted successfully"
    }


@router.get("/")
async def list_documents(db: AsyncSession = Depends(get_db)):
    """List all uploaded documents"""
    result = await db.execute(select(Document).order_by(Document.created_at.desc()))
    documents = result.scalars().all()
    return [{"id": d.id, "filename": d.original_filename, "status": d.status, "created_at": d.created_at} for d in documents]


@router.get("/{document_id}")
async def get_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """Get document details"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document


@router.delete("/{document_id}")
async def delete_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """Delete a document"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file
    if os.path.exists(document.file_path):
        os.remove(document.file_path)
    
    await db.delete(document)
    await db.commit()
    return {"message": "Document deleted successfully"}
